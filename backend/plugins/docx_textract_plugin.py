import docx
from typing import List, Dict, Any
from pathlib import Path
import logging

from infrastructure.plugins.base_plugin import IDataSourcePlugin

logger = logging.getLogger(__name__)


class DOCXTextractPlugin(IDataSourcePlugin):
    """DOCX parser using python-docx for Word document processing."""
    
    @staticmethod
    def get_name() -> str:
        return "docx_textract_parser"
    
    @staticmethod
    def get_supported_identifiers() -> List[str]:
        return [".docx", ".doc"]
    
    async def process(self, source_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX files."""
        
        if not Path(source_path).exists():
            raise FileNotFoundError(f"File not found: {source_path}")
        
        chunks = []
        
        try:
            doc = docx.Document(source_path)
            
            current_page_text = []
            estimated_page_num = 1
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_page_text.append(text)
                    
                    if len(current_page_text) >= 30:
                        chunks.append({
                            "text_content": "\n".join(current_page_text),
                            "metadata": {
                                "page_number": estimated_page_num,
                                "source_file": Path(source_path).name,
                                "parser": self.get_name(),
                            }
                        })
                        current_page_text = []
                        estimated_page_num += 1
            
            if current_page_text:
                chunks.append({
                    "text_content": "\n".join(current_page_text),
                    "metadata": {
                        "page_number": estimated_page_num,
                        "source_file": Path(source_path).name,
                        "parser": self.get_name(),
                    }
                })
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                
                if table_text:
                    chunks.append({
                        "text_content": "\n".join(table_text),
                        "metadata": {
                            "content_type": "table",
                            "source_file": Path(source_path).name,
                            "parser": self.get_name(),
                        }
                    })
            
            logger.info(f"Extracted {len(chunks)} chunks from {source_path}")
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise
        
        return chunks